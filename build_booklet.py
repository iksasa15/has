#!/usr/bin/env python3
"""Generate بوصلة UT printable Word booklet template."""

from __future__ import annotations

import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "بوصلة-UT-قالب.docx"
ASSETS = ROOT / "_assets"

GREEN = "3B5249"
CREAM = "EDE7D4"
PURPLE = "7473B3"
GOLD = "F8B624"
PASTEL_BLUE = "B5C9E4"
MAGENTA = "FF5CB6"
CYAN = "519CF2"
PLUM = "651853"
LEAF = "5BB34E"
BLACK = "000000"
WHITE = "FFFFFF"
GRAY_LINE = "D0D0D0"
CREAM_LINE = "C5BFA8"

# Exact family names from fonts/*.otf (must match Word / macOS)
FONT_REGULAR = "thmanyah sans"
FONT_BOLD = "thmanyah sans"
FONT_MEDIUM = "thmanyah sans Med"
FONT_LIGHT = "thmanyah sans Light"
FONT_BLACK = "thmanyah sans Black"
FONT = FONT_REGULAR  # default body face

A5_W = Cm(14.8)
A5_H = Cm(21.0)


def hex_rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_run_font(
    run,
    size_pt: float,
    bold: bool = False,
    color: str = BLACK,
    name: str | None = None,
    weight: str = "regular",
):
    """Apply thmanyah sans face. weight: regular|bold|medium|light|black."""
    weight_map = {
        "regular": (FONT_REGULAR, False),
        "bold": (FONT_BOLD, True),
        "medium": (FONT_MEDIUM, False),
        "light": (FONT_LIGHT, False),
        "black": (FONT_BLACK, False),
    }
    if name is None:
        name, weight_bold = weight_map.get(weight, (FONT_REGULAR, False))
        bold = bold or weight_bold
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = hex_rgb(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz_cs)
    rPr.append(OxmlElement("w:rtl"))


def set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=0, line=1.2):
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    pPr = p._element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def shade_paragraph(p, fill: str):
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def set_cell_shading(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tcPr.append(tc_mar)


def set_cell_width(cell, cm: float):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    twips = int(cm * 567)
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    cell.width = Cm(cm)


def clear_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tcPr.append(borders)


def set_cell_border(cell, **edges):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        spec = edges.get(edge)
        if spec is None:
            el.set(qn("w:val"), "nil")
        else:
            style, sz, color = spec
            el.set(qn("w:val"), style)
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def clear_paragraph(p):
    for child in list(p._element):
        p._element.remove(child)


def set_section_a5(section, margins=(1.5, 1.5, 1.5, 1.5)):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = A5_W
    section.page_height = A5_H
    top, bottom, left, right = margins
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    section._sectPr.append(bidi)


def new_section(doc: Document, margins=(1.5, 1.5, 1.5, 1.5)):
    """Start a new page via section break (no extra page-break para)."""
    section = doc.add_section()
    set_section_a5(section, margins=margins)
    return section


def thin_accent_bar(container, color: str = GOLD, width_hint_cm: float | None = None):
    tbl = container.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, color)
    clear_cell_borders(cell)
    set_cell_margins(cell, 0, 0, 0, 0)
    if width_hint_cm:
        set_cell_width(cell, width_hint_cm)
    p = cell.paragraphs[0]
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    set_run_font(p.add_run(" "), 2, color=color)
    return tbl


def make_star_png(path: Path, size: int = 128, color: str = GOLD):
    img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    d = ImageDraw.Draw(img)
    cx = cy = size / 2
    rgb = tuple(int(color[i : i + 2], 16) for i in (0, 2, 4))
    pts = []
    for i in range(8):
        ang = math.radians(i * 45 - 90)
        r = size * 0.45 if i % 2 == 0 else size * 0.18
        pts.append((cx + r * math.cos(ang), cy + r * math.sin(ang)))
    d.polygon(pts, fill=rgb + (255,))
    w = max(2, size // 16)
    d.rectangle([cx - w, size * 0.08, cx + w, size * 0.92], fill=rgb + (255,))
    d.rectangle([size * 0.08, cy - w, size * 0.92, cy + w], fill=rgb + (255,))
    img.save(path)


def make_quote_mark_png(path: Path, size: int = 200, color: str = PURPLE):
    img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    d = ImageDraw.Draw(img)
    rgb = tuple(int(color[i : i + 2], 16) for i in (0, 2, 4))
    font = None
    for candidate in (
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
    ):
        if Path(candidate).exists():
            try:
                font = ImageFont.truetype(candidate, int(size * 0.75))
                break
            except Exception:
                continue
    if font:
        d.text((size * 0.05, size * -0.08), "”", fill=rgb + (255,), font=font)
    else:
        d.ellipse([20, 30, 90, 100], fill=rgb + (255,))
        d.ellipse([110, 30, 180, 100], fill=rgb + (255,))
        d.polygon([(20, 100), (55, 170), (90, 100)], fill=rgb + (255,))
        d.polygon([(110, 100), (145, 170), (180, 100)], fill=rgb + (255,))
    img.save(path)


def make_logo_placeholder(path: Path, size: int = 220):
    img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    d = ImageDraw.Draw(img)
    margin = 12
    d.rectangle([margin, margin, size - margin, size - margin], outline=(255, 255, 255, 255), width=3)
    d.arc(
        [size * 0.2, size * 0.25, size * 0.8, size * 0.85],
        start=200,
        end=340,
        fill=(255, 255, 255, 255),
        width=4,
    )
    cx, cy = size / 2, size * 0.55
    for i in range(-4, 5):
        ang = math.radians(-90 + i * 12)
        x2 = cx + math.cos(ang) * size * 0.32
        y2 = cy + math.sin(ang) * size * 0.32
        d.line([(cx, cy - 10), (x2, y2)], fill=(255, 255, 255, 220), width=2)
    img.save(path)


def _apply_style_font(style, family: str, size: float, bold: bool, color: str):
    style.font.name = family
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = hex_rgb(color)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), family)


def define_styles(doc: Document):
    styles = doc.styles
    _apply_style_font(styles["Normal"], FONT_REGULAR, 11, False, BLACK)

    def ensure(name):
        try:
            return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            return styles[name]

    # name, size, bold, color, family
    for name, size, bold, color, family in [
        ("بوصلة-عنوان-رئيسي", 28, False, BLACK, FONT_BLACK),
        ("بوصلة-عنوان-قسم", 16, True, GREEN, FONT_BOLD),
        ("بوصلة-نص-جسم", 10, False, BLACK, FONT_REGULAR),
        ("بوصلة-تسمية-رقمية", 13, False, BLACK, FONT_MEDIUM),
        ("بوصلة-اقتباس", 12, False, PURPLE, FONT_MEDIUM),
        ("بوصلة-فوتر", 8, False, WHITE, FONT_LIGHT),
        ("بوصلة-عنوان-أبيض", 18, True, WHITE, FONT_BOLD),
        ("بوصلة-نص-أبيض", 11, False, WHITE, FONT_REGULAR),
    ]:
        _apply_style_font(ensure(name), family, size, bold, color)


def build_cover(doc: Document):
    meta = doc.add_table(rows=1, cols=2)
    left, right = meta.rows[0].cells
    clear_cell_borders(left)
    clear_cell_borders(right)
    set_cell_width(left, 6.0)
    set_cell_width(right, 6.0)

    p = left.paragraphs[0]
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)
    set_run_font(p.add_run("Aug. 23, 2026"), 9, color=BLACK, weight="light")

    p = right.paragraphs[0]
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    set_run_font(p.add_run("الطبعة الأولى"), 9, color=BLACK, weight="light")

    for _ in range(5):
        sp = doc.add_paragraph()
        set_para_rtl(sp, space_after=0)
        sp.paragraph_format.space_before = Pt(8)

    title = doc.add_paragraph(style="بوصلة-عنوان-رئيسي")
    set_para_rtl(title, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, space_before=8)
    clear_paragraph(title)
    set_run_font(title.add_run("بوصلة UT"), 36, color=BLACK, weight="black")

    sub = doc.add_paragraph()
    set_para_rtl(sub, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    set_run_font(sub.add_run("مذكرة مستجد"), 14, color=BLACK, weight="medium")

    # Centered short gold rule via 3-col table
    rule = doc.add_table(rows=1, cols=3)
    for i, cell in enumerate(rule.rows[0].cells):
        clear_cell_borders(cell)
        set_cell_margins(cell, 0, 0, 0, 0)
        if i == 1:
            set_cell_shading(cell, GOLD)
            set_cell_width(cell, 3.2)
            p = cell.paragraphs[0]
            set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
            set_run_font(p.add_run(" "), 3, color=GOLD)
        else:
            set_cell_width(cell, 4.5)
            p = cell.paragraphs[0]
            set_para_rtl(p, space_after=0)
            set_run_font(p.add_run(" "), 3, color=WHITE)

    for _ in range(10):
        sp = doc.add_paragraph()
        set_para_rtl(sp, space_after=0)
        sp.paragraph_format.space_before = Pt(6)


def build_name_page(doc: Document):
    new_section(doc, margins=(1.2, 1.2, 1.2, 1.2))

    outer = doc.add_table(rows=1, cols=1)
    oc = outer.rows[0].cells[0]
    set_cell_shading(oc, PURPLE)
    clear_cell_borders(oc)
    set_cell_margins(oc, top=280, bottom=240, left=180, right=180)

    clear_paragraph(oc.paragraphs[0])

    # Name field
    name_tbl = oc.add_table(rows=1, cols=2)
    c_line, c_label = name_tbl.rows[0].cells
    for c in (c_line, c_label):
        set_cell_shading(c, PURPLE)
        clear_cell_borders(c)
    set_cell_width(c_line, 7.5)
    set_cell_width(c_label, 3.5)
    set_cell_border(c_line, bottom=("single", "12", WHITE))

    p = c_line.paragraphs[0]
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=0)
    set_run_font(p.add_run(" "), 12, color=WHITE)

    p = c_label.paragraphs[0]
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=6, space_after=0)
    set_run_font(p.add_run("اسمك الكريم"), 11, color=WHITE, weight="medium")

    for _ in range(6):
        sp = oc.add_paragraph()
        set_para_rtl(sp, space_after=0)
        shade_paragraph(sp, PURPLE)
        set_run_font(sp.add_run(" "), 12, color=PURPLE)

    msg1 = oc.add_paragraph(style="بوصلة-عنوان-أبيض")
    set_para_rtl(msg1, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    shade_paragraph(msg1, PURPLE)
    clear_paragraph(msg1)
    set_run_font(msg1.add_run("صممنا هذي النسخة الاستثنائية لك .. حتى"), 12.5, color=WHITE, weight="bold")

    msg2 = oc.add_paragraph(style="بوصلة-عنوان-أبيض")
    set_para_rtl(msg2, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    shade_paragraph(msg2, PURPLE)
    clear_paragraph(msg2)
    set_run_font(msg2.add_run("تكون دليلك الأول"), 12.5, color=WHITE, weight="bold")

    for _ in range(7):
        sp = oc.add_paragraph()
        set_para_rtl(sp, space_after=0)
        shade_paragraph(sp, PURPLE)
        set_run_font(sp.add_run(" "), 12, color=PURPLE)

    sig = oc.add_table(rows=2, cols=2)
    for row in sig.rows:
        for cell in row.cells:
            set_cell_shading(cell, PURPLE)
            clear_cell_borders(cell)
            set_cell_width(cell, 5.5)

    p = sig.rows[0].cells[0].paragraphs[0]
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    set_run_font(p.add_run("توقيعك"), 10, color=WHITE, weight="light")

    p = sig.rows[0].cells[1].paragraphs[0]
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    set_run_font(p.add_run("التاريخ"), 10, color=WHITE, weight="light")

    for cell in sig.rows[1].cells:
        set_cell_border(cell, bottom=("single", "12", WHITE))
        p = cell.paragraphs[0]
        set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=16, space_after=0)
        set_run_font(p.add_run(" "), 10, color=WHITE)

    # Extra purple padding at bottom of outer cell
    for _ in range(2):
        sp = oc.add_paragraph()
        set_para_rtl(sp, space_after=0)
        shade_paragraph(sp, PURPLE)
        set_run_font(sp.add_run(" "), 10, color=PURPLE)


def tip_block(cell, number: str, title: str, body: str):
    set_cell_shading(cell, WHITE)
    clear_cell_borders(cell)
    set_cell_margins(cell, top=50, bottom=70, left=50, right=50)
    p = cell.paragraphs[0]
    clear_paragraph(p)
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)
    set_run_font(p.add_run(f"{number}  {title}"), 10.5, color=BLACK, weight="bold")

    bp = cell.add_paragraph(style="بوصلة-نص-جسم")
    set_para_rtl(bp, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0, line=1.35)
    clear_paragraph(bp)
    set_run_font(bp.add_run(body), 8.5, color=BLACK, weight="regular")


def quote_box_cell(cell, quote_img: Path):
    set_cell_shading(cell, WHITE)
    set_cell_margins(cell, top=70, bottom=70, left=70, right=70)
    set_cell_border(
        cell,
        top=("single", "12", GRAY_LINE),
        bottom=("single", "12", GRAY_LINE),
        left=("single", "12", GRAY_LINE),
        right=("single", "12", GRAY_LINE),
    )
    p = cell.paragraphs[0]
    clear_paragraph(p)
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    p.add_run().add_picture(str(quote_img), width=Cm(1.3))

    bp = cell.add_paragraph(style="بوصلة-اقتباس")
    set_para_rtl(bp, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0, line=1.3)
    clear_paragraph(bp)
    set_run_font(bp.add_run("مساحة لاقتباسك المفضّل أو ملاحظة شخصية…"), 9, color=PURPLE, weight="medium")


def build_tips_page(doc: Document, quote_img: Path):
    new_section(doc, margins=(1.3, 1.3, 1.2, 1.2))

    header = doc.add_table(rows=1, cols=2)
    text_cell, bar = header.rows[0].cells  # LTR: bar on visual right for Arabic feel
    set_cell_width(bar, 0.12)
    set_cell_width(text_cell, 12.0)
    set_cell_shading(bar, BLACK)
    clear_cell_borders(bar)
    clear_cell_borders(text_cell)
    set_cell_margins(bar, 0, 0, 0, 0)
    set_cell_margins(text_cell, top=20, bottom=20, left=40, right=80)

    p = bar.paragraphs[0]
    set_para_rtl(p, space_after=0)
    set_run_font(p.add_run(" "), 18, color=BLACK)

    p = text_cell.paragraphs[0]
    clear_paragraph(p)
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0, line=1.25)
    set_run_font(p.add_run("البدايات لا تُكتب بالحظ... بل بالعادات"), 10, color=BLACK, weight="medium")
    p2 = text_cell.add_paragraph()
    set_para_rtl(p2, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0, line=1.25)
    set_run_font(p2.add_run("الصغيرة التي نختارها كل يوم"), 10, color=BLACK, weight="medium")

    sp = doc.add_paragraph()
    set_para_rtl(sp, space_before=8, space_after=4)
    set_run_font(sp.add_run(" "), 4)

    tips = [
        (
            "01",
            "لا تؤجل البداية",
            "قد يبدو الأسبوع الأول خفيفًا لكنه يرسم الفصل كاملاً. وكل مهمة تنجزها اليوم، هي عبء أقل تحمله غدًا.",
        ),
        (
            "04",
            "عش التجربة كاملة",
            "درجاتك جزء من رحلتك لكنها ليست الرحلة كلها. شارك، تعلّم، واصنع لنفسك ذكريات تستحق أن تُروى بعد التخرج.",
        ),
        (
            "02",
            "اسأل بثقة",
            "ليس السؤال دليلًا على الجهل، بل رغبة في الفهم. وكل إجابة تحصل عليها اليوم، تختصر عليك كثيرًا من التردد.",
        ),
        (
            "05",
            "نظّم وقتك قبل أن ينظّمك الوقت",
            "المهام الجامعية لا تتراكم فجأة بل تتسلل بهدوء. وما تكتبه في جدولك اليوم، لن يطاردك غدًا.",
        ),
        (
            "03",
            "اجعل للمراجعة موعدًا",
            "لا تجعل كتبك لا تعرفك إلا ليلة الاختبار.. دقائق قليلة بعد كل محاضرة قد تغنيك عن ساعات طويلة من القلق.",
        ),
    ]

    grid = doc.add_table(rows=3, cols=2)
    # LTR table: col0=left, col1=right → put odd tips on right (col1)
    tip_block(grid.rows[0].cells[1], *tips[0])
    tip_block(grid.rows[0].cells[0], *tips[1])
    tip_block(grid.rows[1].cells[1], *tips[2])
    tip_block(grid.rows[1].cells[0], *tips[3])
    tip_block(grid.rows[2].cells[1], *tips[4])
    quote_box_cell(grid.rows[2].cells[0], quote_img)


def add_arshidni_footer_table(doc: Document, star_img: Path, logo_img: Path):
    footer = doc.add_table(rows=1, cols=5)
    cells = footer.rows[0].cells
    widths = [1.6, 3.4, 2.0, 3.4, 1.6]
    for c, w in zip(cells, widths):
        set_cell_shading(c, PURPLE)
        clear_cell_borders(c)
        set_cell_margins(c, top=70, bottom=70, left=30, right=30)
        set_cell_width(c, w)

    p = cells[0].paragraphs[0]
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    p.add_run().add_picture(str(star_img), width=Cm(0.42))
    set_run_font(p.add_run(" "), 3, color=WHITE)
    p.add_run().add_picture(str(star_img), width=Cm(0.26))

    p = cells[1].paragraphs[0]
    clear_paragraph(p)
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line=1.1)
    set_run_font(p.add_run("عمادة شؤون الطلاب"), 7, color=WHITE, weight="bold")
    p2 = cells[1].add_paragraph()
    set_para_rtl(p2, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    set_run_font(p2.add_run("Deanship of Students' Affairs"), 5.5, color=WHITE, weight="light")

    p = cells[2].paragraphs[0]
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    p.add_run().add_picture(str(logo_img), width=Cm(1.45))

    p = cells[3].paragraphs[0]
    clear_paragraph(p)
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line=1.1)
    set_run_font(p.add_run("كتيب أرشدني"), 9, color=WHITE, weight="bold")
    p2 = cells[3].add_paragraph()
    set_para_rtl(p2, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    set_run_font(p2.add_run("نسخة 48"), 7, color=WHITE, weight="light")

    p = cells[4].paragraphs[0]
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    for w in (0.26, 0.38, 0.24):
        p.add_run().add_picture(str(star_img), width=Cm(w))
        set_run_font(p.add_run(" "), 2, color=WHITE)


def build_inner_pages(doc: Document, star_img: Path, logo_img: Path):
    new_section(doc, margins=(1.3, 1.0, 1.2, 1.2))

    wrap = doc.add_table(rows=1, cols=1)
    wc = wrap.rows[0].cells[0]
    set_cell_shading(wc, CREAM)
    clear_cell_borders(wc)
    set_cell_margins(wc, top=140, bottom=100, left=110, right=110)

    p = wc.paragraphs[0]
    clear_paragraph(p)
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    set_run_font(p.add_run("محتوى القسم"), 18, color=GREEN, weight="black")

    accent = wc.add_table(rows=1, cols=3)
    for i, cell in enumerate(accent.rows[0].cells):
        clear_cell_borders(cell)
        set_cell_margins(cell, 0, 0, 0, 0)
        set_cell_shading(cell, CREAM if i != 2 else GOLD)
        if i == 2:
            set_cell_width(cell, 2.8)
            set_cell_shading(cell, GOLD)
        p = cell.paragraphs[0]
        set_para_rtl(p, space_after=0)
        set_run_font(p.add_run(" "), 3, color=GOLD if i == 2 else CREAM)

    sp = wc.add_paragraph()
    set_para_rtl(sp, space_before=10, space_after=6)
    shade_paragraph(sp, CREAM)
    set_run_font(sp.add_run(" "), 4, color=CREAM)

    body = wc.add_paragraph(style="بوصلة-نص-جسم")
    set_para_rtl(body, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12, line=1.45)
    shade_paragraph(body, CREAM)
    clear_paragraph(body)
    set_run_font(
        body.add_run(
            "هذه صفحة محتوى داخلية جاهزة للتعديل. استخدم أنماط «بوصلة-عنوان-قسم» و«بوصلة-نص-جسم» "
            "لإضافة فقراتك. الخلفية الكريمية والفوتر البنفسجي يعكسان هوية كتيب أرشدني."
        ),
        10,
        color=GREEN,
        weight="regular",
    )

    chips = wc.add_table(rows=2, cols=4)
    primary = [(GREEN, "أخضر"), (CREAM, "كريمي"), (PURPLE, "بنفسجي"), (GOLD, "ذهبي")]
    secondary = [(PASTEL_BLUE, "سماوي"), (MAGENTA, "وردي"), (CYAN, "أزرق"), (PLUM, "برقوقي")]
    for row, pairs in ((0, primary), (1, secondary)):
        for cell, (col, label) in zip(chips.rows[row].cells, pairs):
            set_cell_shading(cell, col)
            clear_cell_borders(cell)
            set_cell_margins(cell, 50, 50, 30, 30)
            p = cell.paragraphs[0]
            set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
            text_color = GREEN if col in (CREAM, GOLD, PASTEL_BLUE) else WHITE
            set_run_font(p.add_run(label), 7.5, color=text_color, weight="bold")

    note = wc.add_paragraph()
    set_para_rtl(note, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=10, space_after=4)
    shade_paragraph(note, CREAM)
    set_run_font(note.add_run("توزيع الألوان — أساسي / فرعي"), 8, color=GREEN, weight="medium")

    for _ in range(3):
        sp = wc.add_paragraph()
        set_para_rtl(sp, space_after=0)
        shade_paragraph(sp, CREAM)
        set_run_font(sp.add_run(" "), 8, color=CREAM)

    sp = doc.add_paragraph()
    set_para_rtl(sp, space_before=2, space_after=2)
    set_run_font(sp.add_run(" "), 2)
    add_arshidni_footer_table(doc, star_img, logo_img)

    # Notes page
    new_section(doc, margins=(1.3, 1.0, 1.2, 1.2))
    wrap2 = doc.add_table(rows=1, cols=1)
    wc2 = wrap2.rows[0].cells[0]
    set_cell_shading(wc2, CREAM)
    clear_cell_borders(wc2)
    set_cell_margins(wc2, top=140, bottom=80, left=110, right=110)

    p = wc2.paragraphs[0]
    clear_paragraph(p)
    set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=14)
    set_run_font(p.add_run("ملاحظاتي"), 16, color=GREEN, weight="black")

    for _ in range(11):
        line = wc2.add_table(rows=1, cols=1)
        lc = line.rows[0].cells[0]
        set_cell_shading(lc, CREAM)
        set_cell_border(lc, bottom=("single", "6", CREAM_LINE))
        set_cell_margins(lc, top=36, bottom=36, left=20, right=20)
        lp = lc.paragraphs[0]
        set_para_rtl(lp, space_after=0)
        set_run_font(lp.add_run(" "), 11, color=CREAM)

    sp = doc.add_paragraph()
    set_para_rtl(sp, space_before=6, space_after=2)
    set_run_font(sp.add_run(" "), 2)
    add_arshidni_footer_table(doc, star_img, logo_img)


def set_theme_colors_hint(doc: Document):
    core = doc.core_properties
    core.title = "بوصلة UT — مذكرة مستجد"
    core.subject = "كتيب أرشدني نسخة 48"
    core.author = "عمادة شؤون الطلاب"
    core.comments = (
        f"Primary: #{GREEN} #{CREAM} #{PURPLE} #{GOLD} | "
        f"Secondary: #{PASTEL_BLUE} #{MAGENTA} #{CYAN} #{PLUM} #{LEAF} | "
        f"Font: {FONT_REGULAR} / Med / Light / Black"
    )


def main():
    ASSETS.mkdir(exist_ok=True)
    star = ASSETS / "star.png"
    quote = ASSETS / "quote.png"
    logo = ASSETS / "logo_placeholder.png"
    make_star_png(star, size=96, color=GOLD)
    make_quote_mark_png(quote, size=220, color=PURPLE)
    make_logo_placeholder(logo, size=220)

    doc = Document()
    define_styles(doc)
    set_theme_colors_hint(doc)
    set_section_a5(doc.sections[0], margins=(1.6, 1.6, 1.6, 1.6))

    # Remove default empty first paragraph content later by overwriting flow
    build_cover(doc)
    build_name_page(doc)
    build_tips_page(doc, quote)
    build_inner_pages(doc, star, logo)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
